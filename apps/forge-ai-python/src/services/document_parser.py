"""Document parsing service."""

import tempfile
from pathlib import Path
from typing import Any

import pypdf
from docx import Document

from config.settings import get_settings


class DocumentParser:
    """Service for parsing various document formats."""

    SUPPORTED_EXTENSIONS = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "txt": "text/plain",
    }

    async def parse(self, file_content: bytes, filename: str) -> dict[str, Any]:
        """Parse document content.

        Args:
            file_content: Raw file bytes
            filename: Original filename

        Returns:
            Dict with text, pages, and metadata

        Raises:
            ValueError: If file type is not supported
        """
        settings = get_settings()
        extension = Path(filename).suffix.lower().lstrip(".")

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}. Supported: {list(self.SUPPORTED_EXTENSIONS.keys())}")

        # Check file size
        if len(file_content) > settings.max_document_size:
            raise ValueError(f"File too large. Max size: {settings.max_document_size} bytes")

        parser_map = {
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "txt": self._parse_txt,
        }

        return await parser_map[extension](file_content)

    async def _parse_pdf(self, file_content: bytes) -> dict[str, Any]:
        """Parse PDF document."""
        text_parts: list[str] = []
        pages = 0

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
            tmp.write(file_content)
            tmp.flush()

            reader = pypdf.PdfReader(tmp.name)
            pages = len(reader.pages)

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

        full_text = "\n\n".join(text_parts)

        return {
            "text": full_text,
            "pages": pages,
            "metadata": {
                "format": "pdf",
                "pages": pages,
                "characters": len(full_text),
            },
        }

    async def _parse_docx(self, file_content: bytes) -> dict[str, Any]:
        """Parse DOCX document."""
        text_parts: list[str] = []

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
            tmp.write(file_content)
            tmp.flush()

            doc = Document(tmp.name)

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

        full_text = "\n\n".join(text_parts)

        return {
            "text": full_text,
            "pages": 1,  # DOCX doesn't have explicit pages
            "metadata": {
                "format": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "characters": len(full_text),
            },
        }

    async def _parse_txt(self, file_content: bytes) -> dict[str, Any]:
        """Parse plain text document."""
        # Try different encodings
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        text = ""

        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if not text:
            raise ValueError("Could not decode text file with any supported encoding")

        return {
            "text": text,
            "pages": 1,
            "metadata": {
                "format": "txt",
                "characters": len(text),
                "lines": len(text.splitlines()),
            },
        }

    def get_supported_types(self) -> list[dict[str, str]]:
        """Get list of supported document types."""
        return [
            {"extension": ext, "mime_type": mime, "description": f"{ext.upper()} document"}
            for ext, mime in self.SUPPORTED_EXTENSIONS.items()
        ]